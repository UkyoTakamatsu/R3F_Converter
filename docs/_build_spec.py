"""
docs/specification.docx を生成するビルドスクリプト。
実行: python docs/_build_spec.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = "docs/specification.docx"

# ---------------------------------------------------------------------------
# ヘルパー
# ---------------------------------------------------------------------------

def set_font(run, name="Meiryo UI", size=10.5, bold=False, color=None, mono=False):
    run.bold = bold
    run.font.size = Pt(size)
    if mono:
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)
    else:
        run.font.name = name
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Meiryo UI"
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D) if level == 1 else RGBColor(0x2E, 0x74, 0xB5)
    return p


def para(doc, text="", size=10.5, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p


def para_mixed(doc, parts, size=10.5):
    """parts: list of (text, bold, mono)"""
    p = doc.add_paragraph()
    for text, bold, mono in parts:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, mono=mono)
    return p


def code_block(doc, text):
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        set_font(run, mono=True)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "F0F0F0")
        p._p.get_or_add_pPr().append(shading)


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # ヘッダ行
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        set_font(run, bold=True, size=10)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "D6E4F0")
        cell._tc.get_or_add_tcPr().append(shading)

    # データ行
    for r_idx, row in enumerate(rows):
        tr = t.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            mono = val.startswith("`") or val.startswith("  ") or val.startswith("float") or val.startswith("int") or val.startswith("list") or val.startswith("bool") or val.startswith("str") or val.startswith("dict") or val.startswith("tuple") or val.startswith("ndarray") or val.startswith("DataFrame") or val.startswith("Path")
            set_font(run, size=10, mono=bool(val and (val[0].islower() and c_idx == 0 and "(" in val)))

    # 列幅
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return t


def divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# ドキュメント生成
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    # ページ余白
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ===== 表紙 =====
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("R3F Converter")
    run.bold = True
    run.font.name = "Meiryo UI"
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("仕様書 / Specification")
    r2.font.name = "Meiryo UI"
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Tektronix RSA .r3f → CSV / Parquet 変換ツール")
    r3.font.name = "Meiryo UI"
    r3.font.size = Pt(11)

    doc.add_paragraph()
    doc.add_page_break()

    # ===== PART 1: 概要・使い方 =====
    heading(doc, "Part 1  概要と使い方", level=1)
    para(doc, "本書はエンジニア・ユーザー双方を対象とした仕様書です。Part 1 では概要と操作方法を、Part 2 では開発者向けの詳細技術仕様を説明します。")

    divider(doc)

    # 1. 概要
    heading(doc, "1. プロジェクト概要", level=2)
    para(doc,
        "R3F Converter は、Tektronix RSA300/500/600 シリーズのスペクトラムアナライザが録画した "
        ".r3f ファイルを Playback モード（デバイス接続不要）で再生し、"
        "DPX（Digital Phosphor Spectrum）データを CSV または Parquet 形式に変換するツールです。"
    )
    para(doc, "GUI モードとコマンドライン（CLI）モードの両方に対応しています。")

    heading(doc, "主な機能", level=3)
    table(doc,
        ["機能", "説明"],
        [
            ["メタデータ表示",      "中心周波数・サンプルレートを読み取り表示"],
            ["DPX 解析・可視化",   "スペクトラム（Spectrum）とウォーターフォール（Waterfall）をグラフ表示"],
            ["CSV エクスポート",   "frequency_hz / timestamp / amplitude_dbm の 3 列 CSV を出力"],
            ["Parquet エクスポート", "同データを Apache Parquet 形式で出力"],
            ["CLI モード",         "--input 引数でバッチ変換を実行"],
        ],
        col_widths=[4, 12],
    )

    # 2. ファイル構成
    heading(doc, "2. ファイル構成と各モジュールの役割", level=2)
    para(doc, "プロジェクトは以下の 5 つの Python モジュールで構成されます。")

    table(doc,
        ["ファイル", "役割"],
        [
            ["main.py",      "エントリーポイント。引数解析、CLI / GUI の振り分け、.env 設定確認を行う"],
            ["config.py",    ".env ファイルの読み書きと、DLL パスを入力する初期設定 GUI ダイアログを提供する"],
            ["rsa_api.py",   "RSA_API.dll を ctypes でラップする PlaybackRSA クラス。DPX・IQ データ取得の低レベル API"],
            ["converter.py", "DPX データを numpy / pandas で処理し、CSV / Parquet に変換するビジネスロジック層"],
            ["ui.py",        "PyQt6 メインウィンドウ、バックグラウンドスレッド（Worker）、matplotlib グラフ描画"],
        ],
        col_widths=[3.5, 12.5],
    )

    heading(doc, "モジュール依存関係", level=3)
    code_block(doc,
        "main.py\n"
        "  ├── config.py\n"
        "  ├── ui.py\n"
        "  │     ├── rsa_api.py\n"
        "  │     └── converter.py\n"
        "  │           └── rsa_api.py\n"
        "  └── rsa_api.py  (CLI モード直接使用)\n"
        "        └── converter.py"
    )

    # 3. システム要件
    heading(doc, "3. システム要件", level=2)
    table(doc,
        ["項目", "要件"],
        [
            ["OS",         "Windows（RSA_API.dll が Windows 専用のため）"],
            ["Python",     "3.9 以上"],
            ["RSA SDK",    "Tektronix RSA API SDK（RSA_API.dll が必要）"],
            ["アーキテクチャ", "x64"],
        ],
        col_widths=[4, 12],
    )

    heading(doc, "依存 Python ライブラリ", level=3)
    table(doc,
        ["ライブラリ", "用途"],
        [
            ["PyQt6",         "GUI フレームワーク"],
            ["matplotlib",    "グラフ描画（QtAgg バックエンド）"],
            ["numpy",         "配列演算・周波数軸生成"],
            ["pandas",        "DataFrame 生成・CSV / Parquet 保存"],
            ["pyarrow",       "Parquet シリアライズ"],
            ["scipy",         "ピーク検出（scipy.signal.find_peaks）"],
            ["python-dotenv", ".env ファイルの読み込み"],
        ],
        col_widths=[4, 12],
    )

    # 4. セットアップ
    heading(doc, "4. セットアップ", level=2)
    heading(doc, "4.1 インストール", level=3)
    code_block(doc,
        "pip install -r requirements.txt"
    )

    heading(doc, "4.2 .env 設定", level=3)
    para(doc, "RSA_API.dll のパスを .env ファイルに記述します。GUI モードでは初回起動時にダイアログが表示されます。")
    code_block(doc,
        "# .env\n"
        'RSA_API_DLL="C:\\Tektronix\\RSA_API\\lib\\x64\\RSA_API.dll"'
    )

    # 5. 使い方
    heading(doc, "5. 使い方", level=2)

    heading(doc, "5.1 GUI モード", level=3)
    code_block(doc, "python main.py")
    para(doc, "起動後は以下の手順で操作します：")
    steps = [
        ("① ファイルを開く",    "Browse... ボタンで .r3f ファイルを選択"),
        ("② メタデータ確認",    "Read ボタンで中心周波数・サンプルレートを表示"),
        ("③ エクスポート",      "Format (CSV/Parquet) と出力先ディレクトリを選択し Export をクリック"),
        ("④ DPX 解析",         "Acquire DPX & Analyze ボタンでスペクトラムとウォーターフォールを表示"),
    ]
    table(doc,
        ["操作", "内容"],
        [[s, d] for s, d in steps],
        col_widths=[4, 12],
    )

    heading(doc, "5.2 CLI モード", level=3)
    table(doc,
        ["コマンド例", "説明"],
        [
            ["python main.py --input sample.r3f",                    "CSV に変換（デフォルト）"],
            ["python main.py --input sample.r3f --format parquet",   "Parquet に変換"],
            ["python main.py --input sample.r3f --output /tmp/out",  "出力先ディレクトリを指定"],
            ["python main.py --input sample.r3f --meta-only",        "メタデータのみ表示して終了"],
        ],
        col_widths=[8.5, 7.5],
    )

    heading(doc, "5.3 出力データ形式", level=3)
    para(doc, "出力ファイル名は <入力ファイルのステム>.csv / .parquet です。")
    table(doc,
        ["列名", "型", "単位", "説明"],
        [
            ["frequency_hz",  "float64", "Hz",  "絶対 RF 周波数"],
            ["timestamp",     "float64", "s",   "DPX API が返す生のタイムスタンプ（正規化なし）"],
            ["amplitude_dbm", "float64", "dBm", "受信電力"],
        ],
        col_widths=[4, 2.5, 2, 7.5],
    )
    para(doc, "行数 = Hi-Res ライン本数 × トレース長（801 ビン）。Hi-Res ラインが 0 本の場合は -120 dBm のフォールバックデータ 801 行を出力。")

    doc.add_page_break()

    # ===== PART 2: 技術仕様 =====
    heading(doc, "Part 2  開発者向け技術仕様", level=1)
    para(doc, "本パートではモジュールごとのクラス・関数仕様、データフロー、エラーハンドリングを説明します。")

    divider(doc)

    # 6. 起動フロー
    heading(doc, "6. 起動フロー", level=2)

    heading(doc, "6.1 GUI モード", level=3)
    code_block(doc,
        "python main.py\n"
        "  └─ main()                             [main.py:93]\n"
        "       ├─ argparse で引数解析\n"
        "       ├─ args.input が None → GUI モード\n"
        "       ├─ is_env_configured()            [config.py:19]\n"
        "       │    └─ 未設定なら show_initial_setup_dialog()  [config.py:42]\n"
        "       │         └─ QDialog で DLL パス入力 → create_env_from_user_input()\n"
        "       └─ gui_main()                     [main.py:75]\n"
        "            └─ run_gui()                 [ui.py:362]\n"
        "                 ├─ QApplication 作成\n"
        "                 ├─ MainWindow() 作成    [ui.py:128]\n"
        "                 └─ app.exec()           イベントループ開始"
    )

    heading(doc, "6.2 CLI モード", level=3)
    code_block(doc,
        "python main.py --input <path> [--format csv|parquet] [--meta-only]\n"
        "  └─ main()\n"
        "       └─ cli_main()                    [main.py:31]\n"
        "            ├─ (--meta-only 時)\n"
        "            │    └─ PlaybackRSA() → open_r3f_file()\n"
        "            │         → get_center_freq() / get_sample_rate() → 表示して終了\n"
        "            └─ (通常変換)\n"
        "                 ├─ convert_r3f_to_parquet()  または\n"
        "                 └─ convert_r3f_to_csv()\n"
        "                      └─ 進捗を # バーでコンソール表示"
    )

    # 7. モジュール別仕様
    heading(doc, "7. モジュール別仕様", level=2)

    # 7.1 main.py
    heading(doc, "7.1  main.py", level=3)
    table(doc,
        ["関数", "戻り値", "説明"],
        [
            ["main()",                  "int", "メイン処理。CLI / GUI の振り分けと .env チェックを行う"],
            ["cli_main(args)",          "int", "CLI モード処理。変換またはメタデータ表示を実行する"],
            ["gui_main()",              "int", "GUI モード処理。run_gui() を呼び出す"],
            ["build_parser()",          "ArgumentParser", "argparse パーサーを構築して返す"],
        ],
        col_widths=[4.5, 3, 8.5],
    )

    table(doc,
        ["引数", "短縮形", "デフォルト", "説明"],
        [
            ["--input",     "-i", "なし",    ".r3f ファイルパス（省略時は GUI 起動）"],
            ["--output",    "-o", "output",  "出力ディレクトリ"],
            ["--format",    "-f", "csv",     "出力形式（csv または parquet）"],
            ["--meta-only", "-m", "False",   "メタデータのみ表示して終了"],
        ],
        col_widths=[3, 2.5, 3, 7.5],
    )

    # 7.2 config.py
    heading(doc, "7.2  config.py", level=3)
    table(doc,
        ["関数", "戻り値", "説明"],
        [
            ["is_env_configured()",                  "bool", ".env が存在し RSA_API_DLL がデフォルト値以外に設定されているか確認する"],
            ["create_env_from_user_input(dll_path)", "bool", "指定パスで .env を生成する。成功時 True"],
            ["show_initial_setup_dialog()",          "bool", "PyQt6 ダイアログで DLL パス入力を促す。OK かつ .env 生成成功時 True"],
        ],
        col_widths=[5, 2, 9],
    )

    # 7.3 rsa_api.py
    heading(doc, "7.3  rsa_api.py", level=3)
    para(doc, "RSA_API.dll を ctypes.WinDLL でロードし、関数ごとに argtypes / restype を設定して呼び出す Playback モード専用ラッパーです。")

    heading(doc, "PlaybackRSA クラス — 公開メソッド", level=3)
    table(doc,
        ["メソッド", "使用 DLL 関数", "説明"],
        [
            ["open_r3f_file(file_path, ...)",   "PLAYBACK_OpenDiskFile",        "r3f ファイルを開く。失敗時は RSAError を送出"],
            ["close()",                         "DEVICE_Stop",                  "セッションを閉じる"],
            ["get_center_freq() → float",       "CONFIG_GetCenterFreq",         "中心周波数 [Hz] を取得する"],
            ["get_reference_level() → float",   "CONFIG_GetReferenceLevel",     "基準レベル [dBm] を取得する（失敗時は 0.0）"],
            ["get_sample_rate() → float",       "IQBLK_GetIQSampleRate",        "サンプルレート [Hz] を取得する"],
            ["set_record_length(length)",       "IQBLK_SetIQRecordLength",      "1 取得あたりのサンプル数を設定する"],
            ["acquire_iq_data(timeout_ms)",     "IQBLK_AcquireIQData / GetIQData", "IQ データを取得し (I リスト, Q リスト) を返す"],
            ["acquire_dpx_data(fspan, ...)",    "DPX_SetParameters / DEVICE_Run / DPX_WaitForDataReady", "DPX を設定・実行してセッション情報辞書を返す"],
            ["get_dpx_hires_lines(trace_pts)",  "DPX_GetSogramHiResLine",       "DPX 高分解能ラインを全取得する（時系列昇順）"],
            ["diagnose_r3f_file(file_path) ＊静的", "（DLL 不要）",            "r3f ファイルのバイナリ構造を診断する"],
        ],
        col_widths=[5, 5, 6],
    )

    heading(doc, "acquire_dpx_data() 戻り値", level=3)
    table(doc,
        ["キー", "型", "説明"],
        [
            ["trace_length", "int",   "トレースの周波数ビン数（デフォルト 801）"],
            ["fspan",        "float", "スパン [Hz]"],
            ["y_top",        "float", "振幅上限 [dBm]（デフォルト 0.0）"],
            ["y_bottom",     "float", "振幅下限 [dBm]（デフォルト -120.0）"],
        ],
        col_widths=[3.5, 2, 10.5],
    )

    heading(doc, "DPX パラメータ決定ロジック", level=3)
    table(doc,
        ["パラメータ", "決定方法"],
        [
            ["fspan",               "min(sample_rate, 40 MHz)（明示指定も可）"],
            ["rbw",                 "DPX_GetRBWRange() の範囲内で fspan / 200 を基準に決定"],
            ["trace_length",        "801 ビン（固定）"],
            ["y_top",               "0.0 dBm（固定）"],
            ["y_bottom",            "-120.0 dBm（固定）"],
            ["time_per_bitmap_line","0.1 s"],
            ["time_resolution",     "0.01 s（最小 1 ms に丸める）"],
        ],
        col_widths=[5, 11],
    )

    # 7.4 converter.py
    heading(doc, "7.4  converter.py", level=3)
    para(doc, "DPX データを numpy で処理し、pandas DataFrame 経由で CSV / Parquet に変換するビジネスロジック層です。")

    table(doc,
        ["関数", "説明"],
        [
            ["watts_to_dbm(watts) → ndarray",          "電力 [W] → [dBm]。ゼロ/負値は 1e-20 にクリップ"],
            ["dpx_spectrum_arrays(dpx_data, cf)",       "スペクトラムトレースから (freqs_hz, power_dbm) を生成"],
            ["dpx_sogram_arrays(dpx_data, cf)",         "ビットマップから (freqs_hz, times_s, power_matrix_dbm) を生成"],
            ["dpx_hires_to_arrays(hires, cf, fspan)",   "Hi-Res ラインから (freqs, times, power_2d) を生成"],
            ["dpx_to_dataframe(dpx_data, cf, hires)",   "DPX データから (frequency_hz, timestamp, amplitude_dbm) DataFrame を生成"],
            ["save_csv(df, out_path) → Path",           "DataFrame を CSV に保存（float_format='%.6f'）"],
            ["save_parquet(df, out_path) → Path",       "DataFrame を Parquet に保存"],
            ["convert_r3f_to_csv(r3f_path, ...) → Path",     "r3f ファイルを DPX 経由で CSV に変換するメイン関数"],
            ["convert_r3f_to_parquet(r3f_path, ...) → Path", "r3f ファイルを DPX 経由で Parquet に変換するメイン関数"],
        ],
        col_widths=[6.5, 9.5],
    )

    heading(doc, "convert_r3f_to_csv / convert_r3f_to_parquet の変換フロー", level=3)
    code_block(doc,
        "PlaybackRSA() → open_r3f_file()\n"
        "    → get_center_freq() / get_sample_rate()\n"
        "    → acquire_dpx_data(fspan)           # DPX 設定・実行・停止\n"
        "    → get_dpx_hires_lines(trace_len)    # Hi-Res タイムライン取得\n"
        "    → dpx_to_dataframe()                # DataFrame 生成\n"
        "    → save_csv() / save_parquet()        # ファイル保存"
    )
    para(doc, "progress_cb(done, total) は 0/3 → 1/3 → 2/3 → 3/3 の 4 ステップで呼び出されます。")

    # 7.5 ui.py
    heading(doc, "7.5  ui.py", level=3)

    heading(doc, "Worker スレッド", level=3)
    table(doc,
        ["クラス", "シグナル", "説明"],
        [
            ["ConvertWorker(QThread)",  "progress(int,int) / finished(str) / error(str)", "convert_r3f_to_csv / convert_r3f_to_parquet を別スレッドで実行"],
            ["AnalysisWorker(QThread)", "finished(freqs, psd, sog, sr, cf) / error(str)",  "DPX 取得・Hi-Res ライン解析・ndarray 生成を別スレッドで実行"],
        ],
        col_widths=[4, 7, 5],
    )
    para(doc, "Worker はエラーを error シグナルで UI スレッドに通知するため、Worker 内例外が UI をクラッシュさせません。")

    heading(doc, "MainWindow GUI 構成", level=3)
    table(doc,
        ["セクション", "コントロール", "処理"],
        [
            ["1. Open R3F File",  "Browse ボタン",                      "_on_open_file() → QFileDialog → self._r3f_path に保存"],
            ["2. Metadata",       "Read ボタン",                        "_on_read_metadata() → PlaybackRSA で CF / SR 取得 → ラベル更新"],
            ["3. Export Settings","Format コンボ / Output Dir / Export","_on_export() → ConvertWorker 起動 → プログレスバー更新"],
            ["4. DPX Analysis",   "Acquire DPX & Analyze ボタン",       "_on_analyze() → AnalysisWorker 起動 → _on_analysis_done() でグラフ描画"],
        ],
        col_widths=[3.5, 4.5, 8],
    )

    heading(doc, "DPX 解析グラフ仕様", level=3)
    table(doc,
        ["グラフ", "X 軸", "Y 軸", "内容"],
        [
            ["左：DPX Spectrum",          "周波数 [MHz]", "電力 [dBm]", "各ビンの最大値。scipy でピーク上位 5 点を赤▲でマーク。中心周波数を赤破線で表示"],
            ["右：Waterfall (Spectrogram)","時間 [ms]",   "周波数 [MHz]","カラーマップ inferno による電力分布（pcolormesh）。中心周波数をシアン破線で表示"],
        ],
        col_widths=[3.5, 2.5, 2.5, 7.5],
    )

    # 8. データフロー
    heading(doc, "8. エクスポートデータフロー（全体）", level=2)
    code_block(doc,
        ".r3f ファイル\n"
        "    │\n"
        "    ▼  PLAYBACK_OpenDiskFile()\n"
        "RSA_API.dll (Playback モード)\n"
        "    │\n"
        "    ▼  DPX_SetParameters / DPX_SetSogramParameters\n"
        "DPX 設定（fspan, RBW, 振幅レンジ）\n"
        "    │\n"
        "    ▼  DPX_Configure / DPX_SetEnable / DEVICE_Run\n"
        "DPX 実行（Sogram + Spectrum を同時収集）\n"
        "    │\n"
        "    ▼  DPX_WaitForDataReady → DPX_FinishFrameBuffer → DEVICE_Stop\n"
        "    │\n"
        "    ▼  DPX_GetSogramHiResLineCountLatest / DPX_GetSogramHiResLine\n"
        "Hi-Res ライン取得 [(power_dbm[801], timestamp), ...] ← 時系列昇順\n"
        "    │\n"
        "    ▼  dpx_to_dataframe()\n"
        "DataFrame (frequency_hz, timestamp, amplitude_dbm)\n"
        "    │  行数 = n_hires_lines × 801\n"
        "    ▼\n"
        "output/<stem>.csv  または  output/<stem>.parquet"
    )

    # 9. エラーハンドリング
    heading(doc, "9. エラーハンドリング", level=2)

    heading(doc, "PLAYBACK_OpenDiskFile エラーコード", level=3)
    table(doc,
        ["エラーコード", "日本語メッセージ"],
        [
            ["1206", "ファイルを開けません（存在・アクセス権・サイズ・形式を確認）"],
            ["1209", "ファイルフォーマットが正しくないか、破損しています"],
            ["1210", "ファイルが見つかりません"],
            ["1211", "ファイルへのアクセス権がありません"],
            ["その他", "不明なエラー（コード N）"],
        ],
        col_widths=[4, 12],
    )
    para(doc, "エラー発生時は diagnose_r3f_file() の診断情報（ファイルサイズ・マジックバイト等）も RSAError に付加します。")

    heading(doc, "Hi-Res ライン未取得時の動作", level=3)
    table(doc,
        ["呼び出し元", "動作"],
        [
            ["converter.py", "y_bottom (-120 dBm) で埋めたフォールバック DataFrame (801 行) を返す"],
            ["AnalysisWorker", "error シグナルを emit して処理を終了する"],
        ],
        col_widths=[4, 12],
    )

    doc.save(OUTPUT)
    print(f"[OK] {OUTPUT} を生成しました")


if __name__ == "__main__":
    build()
